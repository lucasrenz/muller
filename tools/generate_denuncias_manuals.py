from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
OUT_DIR = BASE_DIR / "docs" / "denuncias"
DOCS_SKILL_SCRIPTS = Path(
    r"C:\Users\Marketing\.codex\plugins\cache\openai-primary-runtime\documents\26.614.11602\skills\documents\scripts"
)

if str(DOCS_SKILL_SCRIPTS) not in sys.path:
    sys.path.insert(0, str(DOCS_SKILL_SCRIPTS))

from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


FONT_BODY = "Calibri"
TITLE_COLOR = RGBColor(0, 0, 0)
HEAD1_COLOR = RGBColor(46, 116, 181)
HEAD2_COLOR = RGBColor(31, 77, 120)
HEAD3_COLOR = RGBColor(31, 77, 120)
MUTED = RGBColor(85, 85, 85)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
TABLE_BORDER = "DADCE0"
NOTE_FILL = "F8FAFC"


def set_run_font(run, *, name=FONT_BODY, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_format(paragraph, *, before=0, after=6, line=1.25, align=None, keep_with_next=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.keep_with_next = keep_with_next
    if align is not None:
      paragraph.alignment = align


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color=TABLE_BORDER, size="8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    core = doc.core_properties
    core.author = "Codex"
    core.company = "Grupo Muller"
    core.comments = "Manual de uso do sistema de denuncias"

    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    h1 = doc.styles["Heading 1"]
    h1.font.name = FONT_BODY
    h1._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = HEAD1_COLOR
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)

    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT_BODY
    h2._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = HEAD2_COLOR
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    h3 = doc.styles["Heading 3"]
    h3.font.name = FONT_BODY
    h3._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
    h3._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = HEAD3_COLOR
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT_BODY
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def add_title_block(doc: Document, kicker: str, title: str, subtitle: str, intro: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(kicker.upper())
    set_run_font(run, size=9.5, color=HEAD1_COLOR, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    set_run_font(run, size=22, color=TITLE_COLOR, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(subtitle)
    set_run_font(run, size=11.5, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(intro)
    set_run_font(run, size=10.5, color=RGBColor(0, 0, 0))


def add_paragraph(doc: Document, text: str, *, before=0, after=6, line=1.25, align=None, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=before, after=after, line=line, align=align)
    run = p.add_run(text)
    set_run_font(run, size=size, color=RGBColor(0, 0, 0), bold=bold, italic=italic)
    return p


def add_mixed_paragraph(doc: Document, runs, *, before=0, after=6, line=1.25, align=None):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=before, after=after, line=line, align=align)
    for text, opts in runs:
        run = p.add_run(text)
        set_run_font(run, **opts)
    return p


def add_heading(doc: Document, text: str, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_bullets(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_format(p, before=0, after=4, line=1.2)
        run = p.add_run(item)
        set_run_font(run, size=11)


def add_numbers(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_paragraph_format(p, before=0, after=4, line=1.2)
        run = p.add_run(item)
        set_run_font(run, size=11)


def add_note(doc: Document, label: str, text: str):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=2, after=8, line=1.2)
    run = p.add_run(f"{label}: ")
    set_run_font(run, size=10.5, color=HEAD2_COLOR, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=RGBColor(0, 0, 0))


def add_table(doc: Document, headers, rows, widths, *, header_fill=LIGHT_BLUE, first_col_fill=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    set_table_borders(table)

    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        set_paragraph_format(p, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run(header)
        set_run_font(run, size=10, bold=True, color=RGBColor(0, 0, 0))

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if col_idx == 0 and first_col_fill:
                set_cell_shading(cell, first_col_fill)
            p = cell.paragraphs[0]
            set_paragraph_format(p, before=0, after=0, line=1.0)
            run = p.add_run(str(value))
            set_run_font(run, size=10.5, color=RGBColor(0, 0, 0), bold=(col_idx == 0 and first_col_fill is not None))

    apply_table_geometry(
        table,
        widths,
        table_width_dxa=sum(widths),
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))

    return table


def blank_line(doc: Document, pts=6):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=0, after=pts, line=1.0)
    return p


def build_user_manual() -> Document:
    doc = Document()
    style_document(doc)
    doc.core_properties.title = "Manual de Utilizacao da Plataforma de Denuncias"
    doc.core_properties.subject = "Manual do usuario do canal de denuncias"
    doc.core_properties.keywords = "denuncias, manual, usuario, consulta"

    add_title_block(
        doc,
        "Canal de Denuncias",
        "Manual de Utilizacao da Plataforma de Denuncias",
        "Fluxo de registro, confirmacao e consulta de denuncias",
        "Documento de apoio para usuarios que precisam registrar uma nova denuncia, acompanhar o protocolo e consultar o andamento com o codigo de acesso.",
    )

    add_table(
        doc,
        ["Item", "Descricao"],
        [
            ("Publico-alvo", "Denunciantes, colaboradores e demais usuarios que acessam o canal publico."),
            ("Abrangencia", "Registro de denuncia, confirmacao de envio e consulta de andamento."),
            ("Acesso", "Rotas publicas /denuncias e /denuncias/consultar, com alias legados /denuncia e /consultar-denuncia."),
            ("Versao", "1.0"),
        ],
        column_widths_from_weights([1.7, 5.8]),
        first_col_fill=LIGHT_GRAY,
    )

    add_heading(doc, "1. Visao geral")
    add_paragraph(
        doc,
        "A plataforma de denuncias foi desenhada para permitir que o usuario relate situacoes de forma segura, com possibilidade de anonimato e consulta posterior pelo protocolo e pelo codigo de acesso gerados no envio.",
    )
    add_paragraph(
        doc,
        "Na pratica, o fluxo tem duas partes: registrar a denuncia e acompanhar o retorno. Em ambos os casos, o usuario trabalha apenas com os dados exibidos na propria interface, sem necessidade de cadastro previo.",
    )

    add_heading(doc, "2. O que pode ser informado")
    add_bullets(
        doc,
        [
            "Assedio moral ou sexual.",
            "Fraude, desvio, corrupcao ou suborno.",
            "Violacao de leis, normas internas ou condutas antieticas.",
        ],
    )
    add_paragraph(
        doc,
        "O formulario tambem aceita detalhes de contexto, como local aproximado, pessoas envolvidas, testemunhas e arquivos de apoio.",
    )

    add_heading(doc, "3. Como registrar uma denuncia")
    add_numbers(
        doc,
        [
            "Acesse a tela de nova denuncia.",
            "Escolha o tipo de ocorrencia.",
            "Informe a data aproximada, o local e a descricao detalhada do fato.",
            "Selecione o vinculo com o Grupo Muller, quando esse dado for relevante.",
            "Inclua pessoas envolvidas e possiveis testemunhas, se houver.",
            "Anexe arquivos de apoio, se desejar.",
            "Escolha se deseja se identificar ou manter anonimato.",
            "Marque a declaracao de veracidade e envie o formulario.",
        ],
    )

    add_note(
        doc,
        "Atencao",
        "O formulario exige tipo de ocorrencia, descricao e aceite dos termos. Se o usuario optar por se identificar, deve preencher ao menos um destes campos: nome, e-mail ou telefone.",
    )

    add_heading(doc, "4. Guia rapido dos campos")
    add_table(
        doc,
        ["Campo", "Como preencher"],
        [
            ("Vinculo com o Grupo Muller", "Escolha colaborador, terceirizado, fornecedor, cliente ou outro."),
            ("Tipo de ocorrencia", "Selecione a categoria que melhor representa o relato."),
            ("Data do ocorrido", "Informe uma data aproximada, se souber."),
            ("Local onde ocorreu", "Descreva a unidade, setor ou ambiente relacionado."),
            ("Descricao", "Explique o que aconteceu, como aconteceu e por que a situacao merece apuracao."),
            ("Pessoas envolvidas", "Liste nomes, cargos ou departamentos, se houver."),
            ("Testemunhas", "Registre pessoas que possam confirmar o fato, se existirem."),
            ("Anexos", "Inclua PDF, JPG, PNG ou MP3 de ate 10 MB por arquivo."),
            ("Identificacao", "Deixe desmarcada para anonimato ou ative a chave para inserir seus dados."),
            ("Termos", "Confirme a declaracao de veracidade antes de enviar."),
        ],
        column_widths_from_weights([1.9, 5.6]),
        first_col_fill=LIGHT_GRAY,
    )

    add_heading(doc, "5. Arquivos e anexos")
    add_bullets(
        doc,
        [
            "Os arquivos podem ser selecionados pelo clique ou arrastados para a area de envio.",
            "Os formatos aceitos sao PDF, JPG, PNG e MP3.",
            "Cada arquivo pode ter ate 10 MB.",
            "Os anexos sao opcionais e servem para reforcar o relato.",
        ],
    )

    add_heading(doc, "6. Identificacao ou anonimato")
    add_bullets(
        doc,
        [
            "Se quiser permanecer anonimo, deixe a opcao de identificacao desativada.",
            "Se quiser informar seus dados, ative a chave e preencha pelo menos um dos campos de contato.",
            "O sistema protege a denuncia em modo anonimo e nao exige login para o envio.",
        ],
    )

    add_heading(doc, "7. Tela de confirmacao")
    add_paragraph(
        doc,
        "Depois do envio bem-sucedido, a plataforma exibe uma tela de confirmacao com o numero do protocolo e o codigo de acesso.",
    )
    add_bullets(
        doc,
        [
            "Guarde os dois dados em local seguro.",
            "O protocolo identifica a denuncia no sistema.",
            "O codigo de acesso nao pode ser recuperado depois.",
            "Esses dados serao usados para consultar o andamento futuramente.",
        ],
    )

    add_heading(doc, "8. Como consultar o andamento")
    add_numbers(
        doc,
        [
            "Acesse a tela de acompanhamento.",
            "Digite o numero do protocolo no formato exibido na tela, por exemplo 2026-1234.",
            "Digite o codigo de acesso de 6 digitos.",
            "Clique em Consultar.",
        ],
    )

    add_heading(doc, "9. O que voce ve no acompanhamento")
    add_bullets(
        doc,
        [
            "O status atual da denuncia.",
            "A data e hora da ultima atualizacao.",
            "O numero do protocolo.",
            "O tipo de ocorrencia registrado.",
            "Uma linha do tempo com as etapas Recebida, Em triagem, Em analise e Concluida.",
            "Mensagens publicas da equipe, quando houver e quando estiverem visiveis ao denunciante.",
            "Uma resposta publica atual, se a equipe tiver publicado um posicionamento oficial.",
        ],
    )

    add_note(
        doc,
        "Importante",
        "No acompanhamento publico, algumas etapas internas sao agrupadas na mesma fase visual. Por isso, estados como em investigacao ou aguardando informacoes continuam aparecendo como Em analise na linha do tempo do denunciante.",
    )

    add_heading(doc, "10. Boas praticas para um bom relato")
    add_bullets(
        doc,
        [
            "Seja objetivo e descreva fatos, datas, locais e pessoas com o maximo de clareza possivel.",
            "Anexe evidencias sempre que isso ajudar na apuracao.",
            "Anote o protocolo e o codigo logo apos o envio.",
            "Use a consulta periodicamente para verificar novidades no caso.",
        ],
    )

    add_heading(doc, "11. Resumo do fluxo")
    add_table(
        doc,
        ["Etapa", "Resultado esperado"],
        [
            ("Registro", "A denuncia e gravada no sistema."),
            ("Confirmacao", "O usuario recebe protocolo e codigo de acesso."),
            ("Acompanhamento", "O usuario consulta o andamento quando quiser."),
            ("Atualizacoes", "Mensagens e respostas publicas podem aparecer conforme a apuracao avanca."),
        ],
        column_widths_from_weights([1.7, 5.9]),
        first_col_fill=LIGHT_GRAY,
    )

    return doc


def build_operator_manual() -> Document:
    doc = Document()
    style_document(doc)
    doc.core_properties.title = "Manual do Operador do Canal de Denuncias"
    doc.core_properties.subject = "Manual operacional do painel juridico"
    doc.core_properties.keywords = "denuncias, operador, juridico, painel, manual"

    add_title_block(
        doc,
        "Guia operacional",
        "Manual do Operador do Canal de Denuncias",
        "Tela de login, painel juridico, fluxo de caso e acoes operacionais",
        "Material de apoio para operadores autorizados que trabalham na triagem, analise, publicacao e acompanhamento das denuncias.",
    )

    add_table(
        doc,
        ["Item", "Descricao"],
        [
            ("Publico-alvo", "Operadores autorizados do modulo de denuncias."),
            ("Abrangencia", "Login, painel, listagem, drawer do caso, notas, anexos, publicacoes e log."),
            ("Acesso", "Rota /denuncias/login com redirecionamento para /denuncias/painel quando a sessao possui permissao."),
            ("Perfil esperado", "Usuario juridico autorizado pelo ambiente de autenticacao."),
        ],
        column_widths_from_weights([1.7, 5.8]),
        first_col_fill=LIGHT_GRAY,
    )

    add_heading(doc, "1. Objetivo do painel")
    add_paragraph(
        doc,
        "O painel juridico concentra a gestao operacional das denuncias. Nele, o operador acompanha a fila, ajusta status e prioridade, registra notas internas, publica respostas e consulta o historico de cada caso.",
    )

    add_heading(doc, "2. Acesso e permissao")
    add_numbers(
        doc,
        [
            "Acesse a tela de login do modulo de denuncias.",
            "Informe e-mail e senha.",
            "Se o perfil tiver permissao, o sistema redireciona automaticamente para o painel.",
            "Se o usuario nao tiver acesso ao modulo, a sessao e encerrada e a tela exibe a mensagem de acesso restrito.",
        ],
    )

    add_note(
        doc,
        "Observacao",
        "A navegacao para o painel e protegida por rota. Em termos praticos, somente usuarios autorizados conseguem acessar /denuncias/painel.",
    )

    add_heading(doc, "3. Visao geral da tela")
    add_table(
        doc,
        ["Area", "O que mostra", "Para que serve"],
        [
            ("Topo", "Nome do operador, busca rapida e botao de sair.", "Controla a sessao e oferece acesso rapido aos comandos principais."),
            ("Cards de resumo", "Total, aguardando triagem, em investigacao e concluidas.", "Mostra o volume do canal e ajuda a priorizar a fila."),
            ("Barra de filtros", "Busca por protocolo ou assunto, status, prioridade e categoria.", "Refina a lista para localizar um caso com rapidez."),
            ("Tabela de denuncias", "Lista paginada com 50 registros por bloco.", "Permite selecionar o caso que sera analisado."),
            ("Resumo lateral", "Descricao, status, prioridade e linha do tempo.", "Mostra a visao rapida do caso selecionado."),
            ("Drawer completo", "Detalhes, notas, anexos, publicacoes e log.", "Concentra as acoes operacionais de um caso."),
        ],
        column_widths_from_weights([1.7, 2.5, 2.3]),
        first_col_fill=LIGHT_GRAY,
    )

    add_heading(doc, "4. Como usar a listagem")
    add_numbers(
        doc,
        [
            "Use os cards de resumo para filtrar rapidamente por status.",
            "Digite no campo de busca para localizar pelo protocolo ou pelo texto da descricao.",
            "Escolha um status, prioridade ou categoria nos filtros da barra superior.",
            "Clique em uma linha da tabela para carregar o resumo do caso na lateral direita.",
            "Clique no botao de seta da linha para abrir o painel completo.",
        ],
    )
    add_bullets(
        doc,
        [
            "A listagem carrega mais registros ao rolar ate o fim da pagina.",
            "A interface mostra 50 casos por bloco durante a navegacao infinita.",
            "No estado atual do sistema, o botao Período funciona como atalho visual e limpa a selecao de status; ele nao abre um seletor de datas.",
        ],
    )

    add_heading(doc, "5. O que aparece no resumo lateral")
    add_bullets(
        doc,
        [
            "Descricao resumida do relato.",
            "Status e prioridade atuais.",
            "Linha do tempo com a denuncia recebida e o status atual.",
            "Atalho para abrir o fluxo completo do caso.",
        ],
    )

    add_heading(doc, "6. Drawer do caso")
    add_paragraph(
        doc,
        "O drawer e o centro de operacao do caso. Ele abre sobre a tela principal e mantem o contexto da denuncia enquanto o operador executa as acoes.",
    )

    add_heading(doc, "6.1 Detalhes")
    add_bullets(
        doc,
        [
            "Mostra tipo, vinculo, data, local, descricao, pessoas envolvidas e testemunhas.",
            "Exibe os dados do denunciante quando a denuncia nao e anonima.",
            "Quando a denuncia e anonima, a interface informa que a identidade nao foi coletada.",
        ],
    )

    add_heading(doc, "6.2 Notas internas")
    add_bullets(
        doc,
        [
            "As notas sao confidenciais e ficam visiveis somente para a equipe interna.",
            "Cada nota mostra autor, setor e data/hora de criacao.",
            "O campo de nova nota fica no rodape da aba e so grava quando ha texto valido.",
        ],
    )

    add_heading(doc, "6.3 Anexos")
    add_bullets(
        doc,
        [
            "A aba lista os arquivos enviados pelo denunciante ou por outros usuarios internos.",
            "Imagens podem ser visualizadas, audios podem ser reproduzidos e PDFs podem ser abertos no proprio painel.",
            "Arquivos sem pre-visualizacao apropriada podem ser abertos em nova aba.",
        ],
    )

    add_heading(doc, "6.4 Publicacoes")
    add_bullets(
        doc,
        [
            "Permite publicar mensagens ao denunciante.",
            "A mensagem pode ter assunto opcional, texto principal e controle de visibilidade.",
            "Se o checkbox de visibilidade estiver marcado, a publicacao aparece na consulta publica do denunciante.",
            "Publicacoes ocultas permanecem apenas no historico interno do caso.",
        ],
    )

    add_heading(doc, "6.5 Log")
    add_bullets(
        doc,
        [
            "Registra as principais atividades da denuncia.",
            "Exibe o tipo de acao, o usuario responsavel, a data e os detalhes do evento.",
            "As mudancas de status e prioridade sao rastreaveis no historico.",
        ],
    )

    add_heading(doc, "7. Alteracao de status e prioridade")
    add_numbers(
        doc,
        [
            "Abra o drawer do caso desejado.",
            "Escolha um novo status ou uma nova prioridade nos selects do topo.",
            "O botao Salvar so aparece quando ha alteracao em relacao aos valores atuais.",
            "Ao salvar, o sistema registra a mudanca e atualiza os dados do caso.",
        ],
    )

    add_table(
        doc,
        ["Status", "Significado operacional"],
        [
            ("nova", "Denuncia recebida e ainda sem triagem inicial."),
            ("triagem", "Triagem em andamento para validar contexto e urgencia."),
            ("em_analise", "Apuracao em curso pelo time responsavel."),
            ("em_investigacao", "Investigacao ativa e com aprofundamento do caso."),
            ("aguardando_informacoes", "Caso parado aguardando retorno, dado ou evidencias."),
            ("concluida", "Apuracao finalizada."),
            ("arquivada", "Caso encerrado sem andamento adicional."),
        ],
        column_widths_from_weights([2.0, 5.5]),
        first_col_fill=LIGHT_GRAY,
    )

    add_heading(doc, "8. Prioridades")
    add_bullets(
        doc,
        [
            "Baixa: demanda sem urgencia operacional.",
            "Normal: acompanhamento padrao.",
            "Alta: caso que pede atencao prioritaria.",
            "Critica: situacao sensivel que exige resposta imediata.",
        ],
    )

    add_heading(doc, "9. Regras de publicacao para o denunciante")
    add_bullets(
        doc,
        [
            "Somente publicacoes marcadas como visiveis aparecem na consulta publica.",
            "A resposta publica atual tambem pode ser exibida no acompanhamento do denunciante.",
            "Notas internas nunca devem ser usadas como mensagem ao denunciante.",
            "Antes de publicar, revise a redacao e confirme se a visibilidade esta correta.",
        ],
    )

    add_heading(doc, "10. Boas praticas operacionais")
    add_bullets(
        doc,
        [
            "Leia a descricao completa antes de alterar status ou prioridade.",
            "Use notas internas para registrar raciocinio, pendencias e decisoes.",
            "Verifique os anexos antes de tomar decisao.",
            "Mantenha o log claro, objetivo e consistente com a evolucao do caso.",
            "Finalize o atendimento com status adequado e, quando fizer sentido, com uma resposta publica adequada ao denunciante.",
        ],
    )

    add_heading(doc, "11. Resumo rapido das acoes")
    add_table(
        doc,
        ["Acao", "Onde fazer", "Efeito"],
        [
            ("Pesquisar caso", "Barra de filtros", "Localiza denuncias por protocolo, descricao, status, prioridade ou categoria."),
            ("Abrir resumo", "Clique na linha", "Carrega a visao rapida do caso na coluna lateral."),
            ("Abrir caso completo", "Botao de seta ou botao Abrir Fluxo Completo", "Exibe o drawer com todas as abas."),
            ("Salvar alteracoes", "Topo do drawer", "Grava status e prioridade quando houve mudanca."),
            ("Adicionar nota", "Aba Notas", "Registra informacao confidencial interna."),
            ("Publicar mensagem", "Aba Publicacoes", "Envia comunicacao ao denunciante, com ou sem visibilidade publica."),
        ],
        column_widths_from_weights([1.9, 2.4, 2.4]),
        first_col_fill=LIGHT_GRAY,
    )

    return doc


def save_doc(doc: Document, path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main():
    user_doc = build_user_manual()
    operator_doc = build_operator_manual()

    user_path = OUT_DIR / "manual_usuario_plataforma_denuncias.docx"
    operator_path = OUT_DIR / "manual_operador_plataforma_denuncias.docx"

    save_doc(user_doc, user_path)
    save_doc(operator_doc, operator_path)

    print(user_path)
    print(operator_path)


if __name__ == "__main__":
    main()
